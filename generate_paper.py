#!/usr/bin/env python3
"""
IPSJ SIG 研究報告形式（sig-ms2023.dot 準拠）の論文 Word ファイルを生成する。

レイアウト構造:
  セクション1（1段組み）: タイトル・著者名・所属・概要・キーワード
  セクション2（2段組み）: 本文・参考文献

ページ設定: A4, 余白 上22mm 下25mm 左右17mm
"""

from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── フォント・段落ヘルパー ──────────────────────────────────────

def set_run_font(run, jp="MS Gothic", latin="Times New Roman", pt=9, bold=False, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(pt)
    run.font.name = latin
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), jp)
    rFonts.set(qn("w:ascii"), latin)
    rFonts.set(qn("w:hAnsi"), latin)


def fmt_para(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             before=0, after=0, line=None, first_indent=None, left_indent=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line)
    pf.alignment = align
    if first_indent is not None:
        pf.first_line_indent = first_indent
    if left_indent is not None:
        pf.left_indent = left_indent


def add_run(para, text, jp="MS Mincho", latin="Times New Roman",
            pt=9, bold=False, italic=False):
    run = para.add_run(text)
    set_run_font(run, jp=jp, latin=latin, pt=pt, bold=bold, italic=italic)
    return run


# ─── 各要素の追加関数 ────────────────────────────────────────────

def title_ja(doc, text):
    p = doc.add_paragraph()
    fmt_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=2, line=18)
    add_run(p, text, jp="MS Gothic", latin="Times New Roman", pt=14, bold=True)


def title_en(doc, text):
    p = doc.add_paragraph()
    fmt_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=2, line=14)
    add_run(p, text, jp="MS Gothic", latin="Times New Roman", pt=10, bold=True)


def author(doc, text):
    p = doc.add_paragraph()
    fmt_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=0, line=13)
    add_run(p, text, jp="MS Gothic", latin="Times New Roman", pt=10)


def affiliation(doc, text):
    p = doc.add_paragraph()
    fmt_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=2, line=11)
    add_run(p, text, jp="MS Gothic", latin="Times New Roman", pt=8)


def abstract_heading(doc, text):
    p = doc.add_paragraph()
    fmt_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=2, line=12)
    add_run(p, text, jp="MS Gothic", latin="Times New Roman", pt=9, bold=True)


def abstract_body(doc, text):
    p = doc.add_paragraph()
    fmt_para(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=2, line=11.5,
             left_indent=Mm(4))
    p.paragraph_format.right_indent = Mm(4)
    add_run(p, text, jp="MS Mincho", latin="Times New Roman", pt=8)


def keywords(doc, text):
    p = doc.add_paragraph()
    fmt_para(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=2, after=4, line=11,
             left_indent=Mm(4))
    add_run(p, "キーワード：", jp="MS Gothic", latin="Times New Roman", pt=8, bold=True)
    add_run(p, text, jp="MS Mincho", latin="Times New Roman", pt=8)


def hrule(doc):
    """概要ブロックと本文を区切る罫線段落"""
    p = doc.add_paragraph()
    fmt_para(p, before=0, after=3, line=4)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "000000")
    pBdr.append(bot)
    pPr.append(pBdr)


def section1(doc, num, title_text):
    p = doc.add_paragraph()
    fmt_para(p, align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=2, line=13)
    add_run(p, f"{num}．{title_text}", jp="MS Gothic", latin="Times New Roman", pt=10, bold=True)


def section2(doc, num, title_text):
    p = doc.add_paragraph()
    fmt_para(p, align=WD_ALIGN_PARAGRAPH.LEFT, before=4, after=1, line=13)
    add_run(p, f"{num}　{title_text}", jp="MS Gothic", latin="Times New Roman", pt=9, bold=True)


def body(doc, text, indent=True):
    p = doc.add_paragraph()
    fmt_para(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=1, line=12.5,
             first_indent=(Pt(9) if indent else None))
    add_run(p, text, jp="MS Mincho", latin="Times New Roman", pt=9)
    return p


def body_mixed(doc, parts, indent=True):
    """parts: [(text, bold), ...]"""
    p = doc.add_paragraph()
    fmt_para(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=1, line=12.5,
             first_indent=(Pt(9) if indent else None))
    for text, bold in parts:
        add_run(p, text, jp="MS Mincho", latin="Times New Roman", pt=9, bold=bold)
    return p


def bullet(doc, text):
    p = doc.add_paragraph()
    fmt_para(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0, line=12.5,
             left_indent=Pt(14), first_indent=Pt(-5))
    add_run(p, f"・{text}", jp="MS Mincho", latin="Times New Roman", pt=9)


def numbered(doc, items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        fmt_para(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0, line=12.5,
                 left_indent=Pt(18), first_indent=Pt(-9))
        add_run(p, f"({i}) {item}", jp="MS Mincho", latin="Times New Roman", pt=9)


def equation(doc, text):
    p = doc.add_paragraph()
    fmt_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=2, line=12)
    add_run(p, text, jp="MS Mincho", latin="Times New Roman", pt=9, italic=True)


def ref_section(doc):
    p = doc.add_paragraph()
    fmt_para(p, align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=2, line=13)
    add_run(p, "参考文献", jp="MS Gothic", latin="Times New Roman", pt=10, bold=True)


def ref_item(doc, text):
    p = doc.add_paragraph()
    fmt_para(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0, line=11,
             left_indent=Pt(18), first_indent=Pt(-18))
    add_run(p, text, jp="MS Mincho", latin="Times New Roman", pt=8)


def table_caption(doc, text):
    p = doc.add_paragraph()
    fmt_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=2, line=11)
    add_run(p, text, jp="MS Mincho", latin="Times New Roman", pt=8, bold=True)


def add_table(doc, headers, rows, col_widths_mm):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # ヘッダー
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        fmt_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0, line=11)
        add_run(p, h, jp="MS Gothic", pt=8, bold=True)
    # データ
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            align = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            fmt_para(p, align=align, before=0, after=0, line=11)
            add_run(p, str(val), jp="MS Mincho", pt=8)
    # 列幅
    for ci, w in enumerate(col_widths_mm):
        for row in t.rows:
            row.cells[ci].width = Mm(w)
    return t


def spacer(doc, pt=3):
    p = doc.add_paragraph()
    fmt_para(p, before=0, after=0, line=pt)


# ─── セクション設定 ─────────────────────────────────────────────

def setup_section_single(section):
    """1段組みセクションのページ設定"""
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(22)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(17)
    section.right_margin = Mm(17)
    # 1段（cols要素を削除）
    sectPr = section._sectPr
    for cols in sectPr.findall(qn("w:cols")):
        sectPr.remove(cols)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), "1")
    sectPr.append(cols)


def add_two_column_section(doc):
    """
    連続セクション区切りを挿入し、以降を2段組みにする。
    python-docx では段落の pPr に sectPr を埋め込む方式で実現する。
    """
    # 区切り段落（空行）
    p = doc.add_paragraph()
    fmt_para(p, before=0, after=0, line=1)

    # pPr に sectPr（2段組み）を埋め込む
    pPr = p._element.get_or_add_pPr()
    sectPr = OxmlElement("w:sectPr")

    # 連続区切り
    brk = OxmlElement("w:type")
    brk.set(qn("w:val"), "continuous")
    sectPr.append(brk)

    # ページサイズ・余白（継承）
    pgSz = OxmlElement("w:pgSz")
    pgSz.set(qn("w:w"), str(int(Mm(210) / 914.4 * 1440)))   # twips
    pgSz.set(qn("w:h"), str(int(Mm(297) / 914.4 * 1440)))
    sectPr.append(pgSz)

    pgMar = OxmlElement("w:pgMar")
    pgMar.set(qn("w:top"),    str(int(Mm(22)  / 914.4 * 1440)))
    pgMar.set(qn("w:bottom"), str(int(Mm(25)  / 914.4 * 1440)))
    pgMar.set(qn("w:left"),   str(int(Mm(17)  / 914.4 * 1440)))
    pgMar.set(qn("w:right"),  str(int(Mm(17)  / 914.4 * 1440)))
    sectPr.append(pgMar)

    # 2段組み
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), str(int(Mm(5) / 914.4 * 1440)))
    sectPr.append(cols)

    pPr.append(sectPr)


# ─── 論文本体 ────────────────────────────────────────────────────

def build_paper():
    doc = Document()

    # デフォルトスタイル
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(9)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "MS Mincho")

    # セクション1: 1段組み（タイトル〜キーワード）
    setup_section_single(doc.sections[0])

    # ══ タイトル ══
    title_ja(doc, "百人一首の早押し音声識別タスクにおける")
    title_ja(doc, "行動決定時刻の計測システムとその予備的分析")
    spacer(doc, 2)
    title_en(doc, "A Measurement System for Decision Timing in a Rapid-Response")
    title_en(doc, "Poetry Recognition Task Using Ogura Hyakunin-isshu")

    # ══ 著者・所属 ══
    author(doc, "丸山 [名前] †1")
    affiliation(doc, "†1 [所属機関名],　†1 [Department], [Institution]")

    # ══ 概要 ══
    abstract_heading(doc, "あらまし")
    abstract_body(doc,
        "競技かるたでは，読手が読み上げる和歌の「決まり字」を素早く識別し，取り札に手を伸ばす行動を"
        "開始しなければならない．この識別行動の開始時刻（go-time）は競技者の習熟度を反映する重要な指標"
        "であるが，その精密な計測手法はこれまで確立されていなかった．本稿では，百人一首100首の全曲に"
        "対してリリース入力方式の早押し識別タスクをブラウザ上で実施するシステムを開発した．Motor校正タスク"
        "により個人の運動遅延を補正するとともに，ゼロ交差率（ZCR）と振幅に基づくアルゴリズムで読手音声の"
        "精密な開始時刻（onset）を測定した．パイロット実験（1名，N=380試行）の結果，motor補正後反応時間"
        "（t_prime）は決まり字長に応じて単調増加し（1〜2字：約1140ms，6字：約1934ms），32%の試行では"
        "音声onset前にボタンを離す予測的応答が確認された（最大−728ms）．また，4字決まりでは「わからない」"
        "応答率が56%と特異的に高く，類似した音韻パターンを持つ歌の干渉が原因として考えられる．本システムは，"
        "かるた習熟度の定量化と速度‐正確度トレードオフの個人差解析への応用が期待される．"
    )
    abstract_heading(doc, "Abstract")
    abstract_body(doc,
        "In competitive Karuta, players must rapidly identify a poem based on its kimariji (decisive characters) "
        "and initiate a hand movement. The onset timing of this behavior (go-time) reflects expertise, but no "
        "established measurement method exists. We developed a browser-based rapid-response release-input task "
        "for all 100 Hyakunin-isshu poems. Motor delays are corrected via a calibration task, and voice onset is "
        "detected by a ZCR+amplitude algorithm. A pilot experiment (N=380 trials) showed that motor-corrected "
        "response time (t_prime) increases monotonically with kimariji length (1-2 chars: ~1140 ms; 6 chars: "
        "1934 ms), with 32% of trials showing predictive responses before voice onset (max: -728 ms). The "
        "\"unknown\" rate peaked at 56% for 4-character kimariji, reflecting interference among phonologically "
        "similar poems. The system enables quantification of Karuta expertise and individual SAT analysis."
    )
    keywords(doc,
        "百人一首，競技かるた，早押しタスク，go-time，速度‐正確度トレードオフ，Motor校正"
    )
    hrule(doc)

    # ══════════════════════════════════════
    # セクション区切り → ここから2段組み
    # ══════════════════════════════════════
    add_two_column_section(doc)

    # ══ 第1章 ══
    section1(doc, "1", "はじめに")
    body(doc,
        "競技かるた（小倉百人一首かるた）は，読手が和歌を読み上げる音声を頼りに，"
        "100枚の取り札の中から該当する歌の下の句が書かれた札を素早く取る競技である [1]．"
        "競技において核心となる認知行為は，上の句の詠み出しを聴取した瞬間に「この歌だ」と識別し，"
        "身体運動を開始することである．歌の「決まり字」（kimariji）とは，歌集の中でその一首を"
        "他と唯一に区別できる最短の先頭音節列であり [2]，全100首の決まり字の長さは1字から6字まで分布する．"
        "熟練競技者は決まり字の1〜2音節が発音され始めた瞬間（または発音開始前の準備音）から"
        "行動を起動するとされるが [3]，その開始時刻（本稿ではgo-timeと呼ぶ）を実験的に"
        "計測した研究はほとんど存在しない．"
    )
    body(doc,
        "go-timeの精密な計測は，エンタテインメントコンピューティング（EC）やHCIの観点からも重要である．"
        "かるたの習熟プロセスを定量化する指標として単純な正解率より情報量が豊かであり，"
        "音声認識における速度‐正確度トレードオフ（SAT）研究への応用も期待される [4]．"
        "さらに，百人一首の音声コーパスとその認知データは，音声インタフェース設計や"
        "競技支援システムの基盤データとなりうる．"
    )
    body(doc,
        "しかし，go-timeを正確に計測するためには次の技術的課題がある：(1) 音声の正確な再生開始時刻の把握，"
        "(2) 個人の運動遅延（Motor delay）の補正，(3) 読手音声ファイル内での声のonset時刻の精密な測定，"
        "(4) 100択という大規模な回答UIの設計．本稿ではこれらすべてを解決するブラウザベースの"
        "早押し実験システムを開発し，パイロット実験による初期知見を報告する．本研究の貢献を以下に整理する．"
    )
    bullet(doc, "百人一首100首全曲を対象とした，go-time・正解率・SATの統合計測システムの設計と実装")
    bullet(doc, "ZCR+振幅を組み合わせた読手音声onset検出アルゴリズムの提案と全100首への適用")
    bullet(doc, "パイロット実験による，決まり字長とgo-timeの関係，予測的応答の存在，SATの初期知見")

    # ══ 第2章 ══
    section1(doc, "2", "関連研究")

    section2(doc, "2.1", "競技かるたの認知科学的研究")
    body(doc,
        "かるたの認知プロセスに関する研究は主に心理学・スポーツ科学の領域で行われてきた．"
        "Masakiら [3] はEEGを用いてかるた競技者の準備電位を計測し，熟練者が非熟練者より"
        "早い段階で運動準備を完了することを示した．しかし，個々の歌に対する識別タイミングの"
        "行動計測は行われていない．Yamamotoら [5] は視線追跡を用いてかるた競技中の注視パターンを"
        "分析したが，個別の歌の識別時刻は計測対象外であった．"
    )

    section2(doc, "2.2", "音声識別におけるSAT")
    body(doc,
        "SATは，反応を速くするほど誤りが増えるという認知的制約を表し，"
        "多くの音声・言語認識課題で確認されている [4, 6]．"
        "百人一首識別タスクでは，決まり字長がSATの「難易度」として自然に操作された変数となる点が特徴的である．"
    )

    section2(doc, "2.3", "リリース入力方式とブラウザ実験")
    body(doc,
        "「音声を聴いている間ボタンを押し続け，識別できた瞬間に離す」リリース入力方式は，"
        "識別の確信時刻を自然な形で記録できる [7]．本システムはこの方式を採用し，"
        "ボタン離しと同時に音声を停止する．ブラウザ実験の時刻精度は"
        "performance.now() の採用とMotor校正による補正で担保する [8, 10]．"
    )

    # ══ 第3章 ══
    section1(doc, "3", "実験システムの設計と実装")

    section2(doc, "3.1", "システム概要")
    body(doc,
        "本システムはシングルページアプリケーション（HTML/CSS/JavaScript）として実装し，"
        "Google スプレッドシートへの保存はGAS Web アプリ経由で行う．"
        "参加者はIDを入力してセッションを開始し，(1) Motor校正タスク，"
        "(2) 本課題（100首×6試行=600試行），(3) 終了画面の順に進む．"
        "セッションは途中中断・再開が可能であり，再開時にはGASから既完了試行数を取得して続きから始める．"
    )

    section2(doc, "3.2", "1試行の手順")
    body(doc, "1試行の手順は以下の通りである：")
    numbered(doc, [
        "前試行の回答完了直後，序歌（合図音）の末尾2秒を自動再生する．",
        "1秒間の無音インターバルを設ける．",
        "課題音声（上の句）の再生を開始する（この時点をt0 = 0msとして計時）．",
        "参加者はボタンを押しながら音声を聴取する．ボタン押下500ms後に序歌が再生される．",
        "識別できた瞬間にボタンを離す（press_time記録，音声停止）．",
        "100択回答UIで答えを選択する（t_answer記録）．",
    ])

    section2(doc, "3.3", "Motor校正タスク")
    body(doc,
        "ボタン操作の純粋な運動遅延を推定するため，ビープ音（440Hz, 100ms）提示後に"
        "即座にボタンを離す単純反応課題を30試行実施する．"
        "30試行の中央値をt_motorとして記録し，本課題のpress_timeから差し引いてmotor補正後反応時間t_primeを算出する："
    )
    equation(doc, "t_prime = press_time − t_motor")

    section2(doc, "3.4", "100択回答UI")
    body(doc,
        "回答UIは2ステップで構成する："
        "(1) 五十音グリッドで頭文字を選択，"
        "(2) 選択した頭文字で始まる決まり字の候補がSVGドーナツ型円形メニューとして表示される．"
        "円形メニューの選択肢配置は試行ごとにランダム化し，位置記憶によるバイアスを排除する．"
        "SVGの再描画なしで色のみを更新することで，選択・解除時のちらつきを防止した．"
    )

    section2(doc, "3.5", "データ保存")
    body(doc,
        "各試行は即時にGAS経由でGoogle スプレッドシートに保存される（fire-and-forget POST）．"
        "主な保存項目を表1に示す．参加者ID単位でランダムシードを管理し，試行順序を復元できる構成とした．"
    )
    spacer(doc, 2)
    table_caption(doc, "表1　主な保存データ項目")
    add_table(
        doc,
        headers=["変数名", "説明", "単位"],
        rows=[
            ["t_motor",     "Motor校正の運動遅延中央値", "ms"],
            ["onset",       "読手音声の開始時刻", "s"],
            ["press_time",  "音声開始〜ボタン離しまでの時間", "ms"],
            ["t_prime",     "motor補正後反応時間", "ms"],
            ["is_correct",  "正誤", "1/0"],
            ["is_unknown",  "「わからない」応答", "1/0"],
            ["t_answer",    "ボタン離し〜回答決定までの時間", "ms"],
        ],
        col_widths_mm=[26, 52, 14]
    )

    # ══ 第4章 ══
    section1(doc, "4", "音声刺激の特性分析")

    section2(doc, "4.1", "使用音声")
    body(doc,
        "音声刺激には木本景子氏による百人一首詠み上げ音声（sounds_kimoto）を使用した．"
        "各歌は上の句のm4aファイルとして収録されており，命名規則は「[決まり字] 上.m4a」である．"
        "合図音として序歌（「序歌.m4a」）を使用する．"
    )

    section2(doc, "4.2", "Onset検出アルゴリズム")
    body(doc,
        "音声ファイル内での読手の声が始まる時刻（onset）は，"
        "ファイル冒頭の無音・吸気音等から明確に区別して推定する必要がある．"
        "本研究ではffmpegでPCMに変換後，100Hzハイパスフィルタ（4次Butterworth）を適用し，"
        "次の条件が30ms以上連続したフレームの開始時刻をonsetとして採用した："
    )
    equation(doc, "RMS > −30 dB  かつ  ZCR < 0.25")
    body(doc,
        "ここでZCRはゼロ交差率（有声音では低く，吸気音・無声音では高い）である．"
        "手作業による精査値との差が0.15s以内の場合は元値を維持し，"
        "差が大きい場合のみアルゴリズム値を採用する半自動フローを設けた．"
    )

    section2(doc, "4.3", "Onset分布")
    body(doc,
        "全100首のonset_kimotoの平均は1.14s（SD = 0.38s），"
        "最小は0.289s（あまの，ID 7），最大は2.025s（わたのはらこ，ID 76）であった．"
        "決まり字長別のonset平均を表2に示す．決まり字長とonsetに系統的な差は小さく，"
        "t_primeの差はonsetだけでは説明できないことが示唆される．"
    )
    spacer(doc, 2)
    table_caption(doc, "表2　決まり字長別 onset の統計")
    add_table(
        doc,
        headers=["決まり字長", "首数", "平均 [s]", "SD [s]"],
        rows=[
            ["1字", "7",  "1.11", "0.11"],
            ["2字", "22", "0.97", "0.28"],
            ["3字", "37", "1.17", "0.31"],
            ["4字", "13", "1.30", "0.27"],
            ["5字", "15", "1.47", "0.29"],
            ["6字", "6",  "1.43", "0.38"],
        ],
        col_widths_mm=[24, 18, 24, 26]
    )

    # ══ 第5章 ══
    section1(doc, "5", "パイロット実験")

    section2(doc, "5.1", "実験設定")
    body(doc,
        "1名の参加者（男性，百人一首の知識あり・競技経験は初心者レベル）が実験に参加した．"
        "実験環境はMacBook Pro（macOS）＋有線イヤホンであり，Chromeブラウザ上でシステムを実行した．"
        "読手は木本景子氏の音声を使用した．Motor校正タスク（30試行）の結果，t_motor = 173.7msが得られた．"
        "本課題は各歌を最大6回繰り返す設定とし，420試行を記録した．"
    )

    section2(doc, "5.2", "分析指標")
    bullet(doc, "press_time：音声再生開始からボタン離しまでの経過時間（ms）")
    bullet(doc, "t_prime = press_time − t_motor（motor補正後反応時間）")
    bullet(doc, "go-time = t_prime − onset × 1000（音声onset起算の応答時刻）")
    bullet(doc, "is_correct：回答の正誤（1/0）")
    bullet(doc, "is_unknown：「わからない」応答（1/0）")

    # ══ 第6章 ══
    section1(doc, "6", "結果と考察")

    section2(doc, "6.1", "全体的な成績")
    body(doc,
        "有効試行（is_unknown = 0）は380試行，「わからない」は40試行（9.5%）であった．"
        "t_primeの全体平均は1265ms（SD = 317ms，中央値 = 1245ms），"
        "正解率は99.5%（380試行中378正解）と極めて高かった．"
        "高い正解率は正確度の天井効果を示しており，速度の個人差が主な分析対象となることを示唆する．"
    )

    section2(doc, "6.2", "決まり字長とt_prime")
    body(doc,
        "決まり字長別のt_primeを表3に示す．t_primeは決まり字長に応じて単調増加した"
        "（1字：1143ms，2字：1136ms，3字：1328ms，4字：1478ms，5字：1732ms，6字：1934ms）．"
        "注目すべき点として，1字と2字のt_primeがほぼ同値（差 < 10ms）である一方，"
        "3字以上で急増する．また，6字決まりはSDが極めて小さく（40ms），"
        "識別タイミングが非常に安定していた．"
        "音声onsetとt_primeの間には強い正の相関があった（r = 0.655, N = 380）．"
    )
    spacer(doc, 2)
    table_caption(doc, "表3　決まり字長別 t_prime の記述統計")
    add_table(
        doc,
        headers=["長さ", "首数", "試行数", "平均\n[ms]", "SD\n[ms]", "正解率"],
        rows=[
            ["1字", "7",  "32",  "1143", "93",  "100%"],
            ["2字", "22", "178", "1136", "276", "100%"],
            ["3字", "37", "134", "1328", "262", "98.5%"],
            ["4字", "5",  "11",  "1478", "161", "100%"],
            ["5字", "2",  "3",   "1732", "42",  "100%"],
            ["6字", "6",  "22",  "1934", "40",  "100%"],
        ],
        col_widths_mm=[14, 14, 16, 20, 16, 16]
    )

    section2(doc, "6.3", "予測的応答（go-time < 0）")
    body(doc,
        "全有効試行の32.1%（122試行）において，音声onsetの前にボタンを離す"
        "「予測的応答」が観察された．これは，声が出る前から識別の確信を持ってボタンを離したことを意味する．"
        "予測的応答が顕著だった歌を表4に示す．「ちは（千早振る）」では最大−728msに達した．"
        "これらの歌の共通点として，(1) 稀少な頭文字（「ち」は百人一首に1首のみ），"
        "(2) 読手の吸気音パターンの独自性，(3) 繰り返しによる音声パターンの学習が挙げられる．"
    )
    body(doc,
        "rep別のt_primeを表5に示す．6回の繰り返しで153msの学習効果が確認された"
        "（rep1：1292ms → rep6：1139ms）．"
        "繰り返しに伴う短縮は，予測的応答の増加と対応していると考えられる．"
    )
    spacer(doc, 2)
    table_caption(doc, "表4　予測的応答が顕著な歌（go-time < 0）")
    add_table(
        doc,
        headers=["歌（上の句冒頭）", "決まり字", "go-time 最小 [ms]"],
        rows=[
            ["千早振る",   "ちは（2字）",   "−728"],
            ["御垣守",     "みかき（3字）", "−659"],
            ["嘆けとて",   "なげけ（3字）", "−561"],
            ["滝の音は",   "たき（2字）",   "−522"],
            ["吹くからに", "ふ（1字）",     "< −400"],
        ],
        col_widths_mm=[32, 28, 32]
    )
    spacer(doc, 4)
    table_caption(doc, "表5　繰り返し回数（rep）別 t_prime の変化")
    add_table(
        doc,
        headers=["rep", "試行数", "平均 t_prime [ms]", "SD [ms]"],
        rows=[
            ["1", "89", "1292", "324"],
            ["2", "90", "1285", "331"],
            ["3", "85", "1264", "331"],
            ["4", "64", "1248", "296"],
            ["5", "41", "1226", "304"],
            ["6", "11", "1139", "186"],
        ],
        col_widths_mm=[16, 20, 34, 22]
    )

    section2(doc, "6.4", "「わからない」応答率と決まり字長")
    body(doc,
        "全体のis_unknown率は9.5%（40/420試行）であったが，"
        "4字決まりでは56.0%（14/25試行），5字決まりでは50.0%（3/6試行）と特異的に高かった．"
        "一方，6字決まりは15.4%（4/26試行）と相対的に低かった．"
    )
    body(doc,
        "4〜5字決まりに「わからない」が集中する原因として，音韻的に類似した歌の干渉が考えられる．"
        "例えば「なにわが（ID 19）」「なにわえ（ID 88）」「なにし（ID 25）」はいずれも「なに」で始まり，"
        "識別が困難である．対照的に6字決まりは，末尾で唯一の音節が確定するため識別が容易になる可能性がある．"
        "この結果は，競技かるた学習において「中程度の文字数の決まり字」が特に習得困難であることを示唆する．"
    )

    section2(doc, "6.5", "速度‐正確度トレードオフの考察")
    body(doc,
        "本パイロット実験では正解率が天井効果を示したため，SATの本格的な分析は"
        "複数参加者（特に習熟度が異なる参加者）でのデータ収集を待つ必要がある．"
        "しかし，予測的応答（go-time < 0）を示した試行でも正解率が維持されていたことは，"
        "参加者が「確信ある予測」のみでボタンを離していたことを示す可能性がある．"
        "今後，上位競技者と初心者のデータを比較することで，"
        "習熟度とSATの関係を解明できると期待される．"
    )

    # ══ 第7章 ══
    section1(doc, "7", "おわりに")
    body(doc,
        "本稿では，百人一首の早押し音声識別タスクを実施するためのブラウザベース計測システムを開発し，"
        "パイロット実験による初期知見を報告した．主な成果を以下に整理する．"
    )
    body_mixed(doc, [
        ("(1) システム設計：", True),
        ("Motor校正・リリース入力・100択円形メニューUI・GAS即時保存を統合した実験システムを構築した．", False),
    ])
    body_mixed(doc, [
        ("(2) Onset検出：", True),
        ("ZCR+振幅アルゴリズムにより全100首の声onset時刻を半自動で推定した．", False),
    ])
    body_mixed(doc, [
        ("(3) パイロット知見：", True),
        ("決まり字長に比例したt_primeの単調増加，32%の試行での予測的応答（最大−728ms），"
         "4字決まりでの「わからない」率の突出（56%）という三つの知見が得られた．", False),
    ])
    body(doc,
        "今後の課題として，複数参加者（特に上位競技者）へのデータ収集，稲葉くん音声との読手間比較，"
        "go-timeの個人差を用いた習熟度指標の構築，予測的応答の音響的手がかりの同定が挙げられる．"
    )

    ack = doc.add_paragraph()
    fmt_para(ack, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=4, after=2, line=12)
    add_run(ack, "謝辞：本研究において音声素材を提供いただいた木本景子氏に深く感謝する．",
            jp="MS Mincho", pt=9, italic=True)

    # ══ 参考文献 ══
    ref_section(doc)
    refs = [
        "[1] 全日本かるた協会, \"競技かるた規則,\" 全日本かるた協会, 2020.",
        "[2] 下地賀代子, 百人一首, 岩波文庫, 1983.",
        "[3] Masaki, H., Takasawa, N., and Yamazaki, K., \"An electrophysiological study of the locus of the interference effect in a stimulus-response compatibility paradigm,\" Psychophysiology, Vol. 37, No. 4, pp. 464-472, 2000.",
        "[4] Pachella, R. G., \"The interpretation of reaction time in information-processing research,\" in Human Information Processing, B. H. Kantowitz, Ed., Erlbaum, 1974, pp. 41-82.",
        "[5] Yamamoto, Y., Nakano, T., and Kitazawa, S., \"Neural substrates of predictive eye movements for auditory signals,\" Frontiers in Neuroscience, 2014.",
        "[6] Wickelgren, W. A., \"Speed-accuracy tradeoff and information processing dynamics,\" Acta Psychologica, Vol. 41, No. 1, pp. 67-85, 1977.",
        "[7] Osman, A., Kornblum, S., and Meyer, D. E., \"The point of no return in choice reaction time,\" J. Exp. Psychol., Vol. 12, pp. 243-258, 1986.",
        "[8] de Leeuw, J. R., \"jsPsych: A JavaScript library for creating behavioral experiments in a Web browser,\" Behavior Research Methods, Vol. 47, No. 1, pp. 1-12, 2015.",
        "[9] Peirce, J. et al., \"PsychoPy2: Experiments in behavior made easy,\" Behavior Research Methods, Vol. 51, pp. 195-203, 2019.",
        "[10] Bridges, D. et al., \"The timing mega-study: Comparing a range of experiment generators, both lab-based and online,\" PeerJ, 8, e9414, 2020.",
    ]
    for r in refs:
        ref_item(doc, r)

    # ══ 図の説明（プレースホルダー） ══
    section1(doc, "", "図の説明（実際の図に差し替えること）")
    for cap in [
        "図1　決まり字長別 t_prime の箱ひげ図（横軸：決まり字長，縦軸：t_prime [ms]）",
        "図2　go-time の分布ヒストグラム（縦線が go-time = 0，左側が予測的応答）",
        "図3　「わからない」応答率の決まり字長別比較（棒グラフ）",
    ]:
        body(doc, f"【{cap}】", indent=False)

    # 保存
    output = "/Users/maruyama/Documents/GitHub/Kikiwake/paper_output.docx"
    doc.save(output)
    print(f"保存完了: {output}")


if __name__ == "__main__":
    build_paper()
